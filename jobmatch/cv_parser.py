from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

from docx import Document
from pypdf import PdfReader


class CVParseError(ValueError):
    """Raised when a CV cannot be converted to useful text."""


MAX_CV_BYTES = 5 * 1024 * 1024
ALLOWED_MIME_TYPES = {
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",
    },
    ".txt": {"text/plain", "application/octet-stream"},
}


def validate_cv_upload(
    filename: str,
    content: bytes,
    mime_type: str | None = None,
) -> None:
    """Reject oversized files and obvious extension/content mismatches."""

    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_MIME_TYPES:
        raise CVParseError("Use a PDF, DOCX, or TXT file.")
    if not content:
        raise CVParseError("The uploaded CV is empty.")
    if len(content) > MAX_CV_BYTES:
        raise CVParseError("The CV is larger than the 5 MB limit.")
    if mime_type and mime_type not in ALLOWED_MIME_TYPES[suffix]:
        raise CVParseError("The file type does not match its filename.")
    if suffix == ".pdf" and not content.startswith(b"%PDF"):
        raise CVParseError("The file is named PDF but does not contain valid PDF data.")
    if suffix == ".docx" and not content.startswith(b"PK"):
        raise CVParseError("The file is named DOCX but does not contain valid DOCX data.")
    if suffix == ".txt" and b"\x00" in content[:4096]:
        raise CVParseError("The TXT upload appears to contain binary data.")


def anonymize_cv_text(text: str, names: tuple[str, ...] = ()) -> str:
    """Remove common personal identifiers before matching, without storing text."""

    redacted = re.sub(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", "[EMAIL]", text, flags=re.I)
    redacted = re.sub(r"https?://\S+|www\.\S+", "[URL]", redacted, flags=re.I)
    redacted = re.sub(
        r"(?<!\w)(?:\+?\d[\d\s()./-]{7,}\d)(?!\w)",
        "[PHONE]",
        redacted,
    )
    for name in sorted((name.strip() for name in names if name.strip()), key=len, reverse=True):
        redacted = re.sub(rf"\b{re.escape(name)}\b", "[NAME]", redacted, flags=re.I)
    return redacted


def extract_cv_text(
    filename: str,
    content: bytes,
    mime_type: str | None = None,
) -> str:
    """Extract text from a PDF, DOCX, or TXT CV without saving it to disk."""
    suffix = Path(filename).suffix.lower()
    validate_cv_upload(filename, content, mime_type=mime_type)

    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
            text = "\n".join(paragraphs + table_cells)
        elif suffix == ".txt":
            text = content.decode("utf-8", errors="replace")
        else:  # guarded by validate_cv_upload; retained for defensive clarity
            raise CVParseError("Use a PDF, DOCX, or TXT file.")
    except CVParseError:
        raise
    except Exception as exc:
        raise CVParseError(f"The CV could not be read: {exc}") from exc

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(text) < 100:
        raise CVParseError(
            "Very little text was extracted. The PDF may be scanned; OCR is not "
            "included in this first version. Try a text-based PDF or DOCX."
        )
    return text
